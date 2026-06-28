from __future__ import annotations

import base64
import hashlib
import hmac
import json
import mimetypes
import os
import re
import secrets
import shutil
import sys
import threading
import time
import traceback
import urllib.parse
from datetime import datetime, timezone, timedelta
from http import HTTPStatus
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path

try:
    from docx import Document
except Exception:  # pragma: no cover - optional runtime dependency
    Document = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas
except Exception:  # pragma: no cover - optional runtime dependency
    A4 = None
    canvas = None
    pdfmetrics = None
    TTFont = None


ROOT = Path(__file__).resolve().parent
DATA_DIR = ROOT / "data"
STORAGE_DIR = ROOT / "storage"
UPLOAD_DIR = STORAGE_DIR / "uploads"
EXPORT_DIR = STORAGE_DIR / "exports"
DB_PATH = DATA_DIR / "db.json"
SECRET = os.environ.get("REPORT_BACKEND_SECRET", "dev-secret-change-me")
TZ = timezone(timedelta(hours=8))


REPORT_TYPES = {
    "summerCheck": "迎峰度夏检查报告",
    "coalInventoryAudit": "煤库存审计报告",
}
REPORT_STATUSES = {
    "draft",
    "outlineGenerated",
    "generating",
    "generated",
    "exporting",
    "exported",
    "generateFailed",
    "exportFailed",
}
CHAPTER_STATUSES = {"pending", "running", "done", "failed"}
ROLES = {"user", "admin", "superAdmin"}
USER_STATUSES = {"enabled", "disabled"}
FILE_FORMATS = {"docx", "pdf", "md", "txt"}


def now_iso() -> str:
    return datetime.now(TZ).replace(microsecond=0).isoformat()


def trace_id() -> str:
    return "trace_" + datetime.now(TZ).strftime("%Y%m%d%H%M%S%f")[:-3]


def ensure_dirs() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)


def hash_password(password: str, salt: str | None = None) -> str:
    salt = salt or secrets.token_hex(16)
    digest = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt.encode("utf-8"), 120_000)
    return f"{salt}${base64.b64encode(digest).decode('ascii')}"


def verify_password(password: str, stored: str) -> bool:
    try:
        salt, digest = stored.split("$", 1)
    except ValueError:
        return False
    expected = hash_password(password, salt).split("$", 1)[1]
    return hmac.compare_digest(expected, digest)


def default_db() -> dict:
    created = now_iso()
    return {
        "counters": {
            "usr": 3,
            "rpt": 0,
            "chap": 0,
            "tbl": 0,
            "tpl": 2,
            "mat": 0,
            "exp": 0,
            "cfg": 1,
        },
        "users": {
            "usr_001": {
                "userId": "usr_001",
                "username": "student",
                "passwordHash": hash_password("123456"),
                "role": "user",
                "displayName": "学生用户",
                "status": "enabled",
                "createdAt": created,
            },
            "usr_002": {
                "userId": "usr_002",
                "username": "admin",
                "passwordHash": hash_password("admin123"),
                "role": "admin",
                "displayName": "管理员",
                "status": "enabled",
                "createdAt": created,
            },
            "usr_003": {
                "userId": "usr_003",
                "username": "super",
                "passwordHash": hash_password("super123"),
                "role": "superAdmin",
                "displayName": "超级管理员",
                "status": "enabled",
                "createdAt": created,
            },
        },
        "sessions": {},
        "reports": {},
        "outlines": {},
        "contents": {},
        "exports": {},
        "templates": {
            "tpl_001": {
                "templateId": "tpl_001",
                "templateName": "迎峰度夏默认模板",
                "reportType": "summerCheck",
                "fileName": "summer-template.docx",
                "filePath": "",
                "status": "enabled",
                "structure": {
                    "titleStyle": "Heading1",
                    "bodyStyle": "Normal",
                    "tableStyle": "GridTable",
                },
                "createdBy": "usr_002",
                "createdAt": created,
            },
            "tpl_002": {
                "templateId": "tpl_002",
                "templateName": "煤库存审计默认模板",
                "reportType": "coalInventoryAudit",
                "fileName": "coal-template.docx",
                "filePath": "",
                "status": "enabled",
                "structure": {
                    "titleStyle": "Heading1",
                    "bodyStyle": "Normal",
                    "tableStyle": "GridTable",
                },
                "createdBy": "usr_002",
                "createdAt": created,
            },
        },
        "materials": {},
        "modelConfig": {
            "configId": "cfg_001",
            "apiUrl": "https://api.example.com/v1/chat/completions",
            "apiKey": "",
            "modelName": "report-model",
            "timeoutSeconds": 120,
            "enabled": True,
            "updatedAt": created,
        },
    }


class Store:
    def __init__(self, path: Path):
        self.path = path
        self.lock = threading.RLock()
        ensure_dirs()
        if not self.path.exists():
            self.data = default_db()
            self.save()
        else:
            self.data = json.loads(self.path.read_text(encoding="utf-8"))
            self._migrate()
            self.save()

    def _migrate(self) -> None:
        base = default_db()
        for key, value in base.items():
            self.data.setdefault(key, value)
        self.data.setdefault("counters", {}).setdefault("usr", len(self.data.get("users", {})))
        self.data.setdefault("exports", {})
        self.data.setdefault("sessions", {})

    def save(self) -> None:
        with self.lock:
            tmp = self.path.with_suffix(".tmp")
            tmp.write_text(json.dumps(self.data, ensure_ascii=False, indent=2), encoding="utf-8")
            tmp.replace(self.path)

    def next_id(self, prefix: str) -> str:
        with self.lock:
            counters = self.data.setdefault("counters", {})
            counters[prefix] = int(counters.get(prefix, 0)) + 1
            return f"{prefix}_{counters[prefix]:03d}"


store = Store(DB_PATH)


def public_user(user: dict) -> dict:
    return {
        "userId": user["userId"],
        "username": user["username"],
        "role": user["role"],
        "displayName": user.get("displayName") or user["username"],
        "createdAt": user.get("createdAt"),
        "status": user.get("status", "enabled"),
    }


def mask_model_config(config: dict) -> dict:
    masked = {k: v for k, v in config.items() if k != "apiKey"}
    api_key = config.get("apiKey") or ""
    masked["apiKeyMasked"] = f"{api_key[:3]}****{api_key[-4:]}" if len(api_key) >= 8 else ""
    return masked


def success(data=None) -> dict:
    return {"code": 0, "message": "success", "data": data, "traceId": trace_id()}


class ApiError(Exception):
    def __init__(self, code: int, message: str, http_status: int = 200, data=None):
        super().__init__(message)
        self.code = code
        self.message = message
        self.http_status = http_status
        self.data = data


def page_items(items: list, page: int, page_size: int) -> dict:
    page = max(int(page or 1), 1)
    page_size = min(max(int(page_size or 10), 1), 100)
    total = len(items)
    total_pages = (total + page_size - 1) // page_size if total else 0
    start = (page - 1) * page_size
    return {
        "items": items[start : start + page_size],
        "page": page,
        "pageSize": page_size,
        "total": total,
        "totalPages": total_pages,
    }


def normalize_outline(report_id: str, outline: list[dict]) -> list[dict]:
    by_temp: dict[str, str] = {}
    normalized: list[dict] = []
    sibling_counts: dict[str | None, int] = {}
    existing = {c["chapterId"]: c for c in store.data.get("outlines", {}).get(report_id, [])}

    for item in outline:
        chapter_id = item.get("chapterId")
        if not chapter_id or chapter_id not in existing:
            chapter_id = store.next_id("chap")
            if item.get("chapterId"):
                by_temp[item["chapterId"]] = chapter_id
        parent_id = item.get("parentId")
        if parent_id in by_temp:
            parent_id = by_temp[parent_id]
        level = int(item.get("level") or (existing.get(chapter_id) or {}).get("level") or 1)
        level = min(max(level, 1), 3)
        title = str(item.get("title") or (existing.get(chapter_id) or {}).get("title") or "未命名章节").strip()
        sibling_counts[parent_id] = sibling_counts.get(parent_id, 0) + 1
        normalized.append(
            {
                "chapterId": chapter_id,
                "reportId": report_id,
                "parentId": parent_id,
                "chapterNo": "",
                "title": title,
                "level": level,
                "sortOrder": int(item.get("sortOrder") or sibling_counts[parent_id]),
                "status": item.get("status") if item.get("status") in CHAPTER_STATUSES else existing.get(chapter_id, {}).get("status", "pending"),
            }
        )

    return renumber_outline(normalized)


def renumber_outline(outline: list[dict]) -> list[dict]:
    children: dict[str | None, list[dict]] = {}
    for chapter in outline:
        children.setdefault(chapter.get("parentId"), []).append(chapter)
    for siblings in children.values():
        siblings.sort(key=lambda c: (int(c.get("sortOrder", 0)), c["chapterId"]))

    ordered: list[dict] = []

    def visit(parent_id: str | None, prefix: str = "") -> None:
        for idx, chapter in enumerate(children.get(parent_id, []), start=1):
            no = f"{prefix}.{idx}" if prefix else str(idx)
            chapter["chapterNo"] = no
            chapter["level"] = no.count(".") + 1
            chapter["sortOrder"] = idx
            ordered.append(chapter)
            visit(chapter["chapterId"], no)

    visit(None)
    # If a malformed parentId hid chapters, keep them as top-level fallbacks.
    seen = {c["chapterId"] for c in ordered}
    for chapter in outline:
        if chapter["chapterId"] not in seen:
            chapter["parentId"] = None
            chapter["level"] = 1
            ordered.append(chapter)
    return ordered


def default_outline(report: dict) -> list[dict]:
    if report["reportType"] == "coalInventoryAudit":
        tree = [
            ("审计概况", ["审计范围", "审计依据", "审计方法"]),
            ("煤库存基础情况", ["库存结构", "入库出库情况", "重点数据说明"]),
            ("审计发现问题", ["账实一致性", "计量与记录", "管理风险"]),
            ("整改建议与结论", ["整改措施", "后续跟踪", "审计结论"]),
        ]
    else:
        tree = [
            ("检查概况", ["检查范围", "检查依据", "检查方法"]),
            ("设备运行情况", ["主设备状态", "辅助系统状态", "隐患排查情况"]),
            ("问题分析与整改建议", ["主要问题", "原因分析", "整改建议"]),
            ("迎峰度夏保障措施", ["组织保障", "技术措施", "应急预案"]),
        ]

    result: list[dict] = []
    for i, (title, children) in enumerate(tree, start=1):
        parent_id = store.next_id("chap")
        result.append(
            {
                "chapterId": parent_id,
                "reportId": report["reportId"],
                "parentId": None,
                "chapterNo": str(i),
                "title": title,
                "level": 1,
                "sortOrder": i,
                "status": "pending",
            }
        )
        for j, child in enumerate(children, start=1):
            result.append(
                {
                    "chapterId": store.next_id("chap"),
                    "reportId": report["reportId"],
                    "parentId": parent_id,
                    "chapterNo": f"{i}.{j}",
                    "title": child,
                    "level": 2,
                    "sortOrder": j,
                    "status": "pending",
                }
            )
    return result


def build_chapter_content(report: dict, chapter: dict) -> tuple[str, list[dict]]:
    report_type_name = REPORT_TYPES.get(report.get("reportType"), "专业报告")
    topic = report.get("topic") or report.get("reportName") or "报告主题"
    base = (
        f"{chapter['chapterNo']} {chapter['title']}：本节围绕“{topic}”展开，结合"
        f"{report.get('plant', '相关单位')}、{report.get('major', '相关专业')}和{report.get('year', '')}年工作要求，"
        f"形成符合{report_type_name}语境的分析内容。"
    )
    detail = (
        "系统依据已确认的大纲、报告类型、模板结构和专业素材组织表述，"
        "重点说明现状、发现、影响和建议，便于后续在线编辑和统一导出。"
    )
    tables: list[dict] = []
    title = chapter.get("title", "")
    if any(key in title for key in ["设备", "问题", "库存", "检查", "审计"]):
        tables.append(
            {
                "tableId": store.next_id("tbl"),
                "title": f"{chapter['title']}情况表",
                "headers": ["序号", "项目", "情况", "建议"],
                "rows": [["1", title, "已形成初步分析", "结合现场数据进一步完善"]],
            }
        )
    return base + "\n\n" + detail, tables


def report_detail(report_id: str) -> dict:
    report = store.data["reports"].get(report_id)
    if not report:
        raise ApiError(40400, "资源不存在")
    outline = store.data.get("outlines", {}).get(report_id, [])
    contents_map = store.data.get("contents", {}).get(report_id, {})
    contents = [contents_map[k] for k in contents_map]
    exports = [e for e in store.data.get("exports", {}).values() if e["reportId"] == report_id]
    exports.sort(key=lambda e: e.get("createdAt", ""), reverse=True)
    return {"report": report, "outline": outline, "contents": contents, "latestExport": exports[0] if exports else None}


def user_can_access_report(user: dict, report: dict) -> bool:
    return user["role"] in {"admin", "superAdmin"} or report.get("createdBy") == user["userId"]


def safe_filename(name: str) -> str:
    cleaned = re.sub(r'[<>:"/\\\\|?*\\x00-\\x1f]', "_", name).strip()
    return cleaned or "file"


def export_markdown(report: dict, outline: list[dict], contents: dict) -> str:
    lines = [f"# {report.get('reportName') or '报告'}", ""]
    for chapter in outline:
        level = max(1, min(int(chapter.get("level", 1)), 6))
        lines.append("#" * level + f" {chapter['chapterNo']} {chapter['title']}")
        content = contents.get(chapter["chapterId"], {})
        if content.get("content"):
            lines.extend(["", content["content"], ""])
        for table in content.get("tables", []):
            headers = table.get("headers") or []
            rows = table.get("rows") or []
            if table.get("title"):
                lines.append(f"**{table['title']}**")
            if headers:
                lines.append("| " + " | ".join(headers) + " |")
                lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
            for row in rows:
                lines.append("| " + " | ".join(map(str, row)) + " |")
            lines.append("")
    return "\n".join(lines)


def write_docx(path: Path, report: dict, outline: list[dict], contents: dict) -> None:
    if Document is None:
        raise RuntimeError("python-docx is not available")
    doc = Document()
    doc.add_heading(report.get("reportName") or "报告", 0)
    doc.add_paragraph(f"报告类型：{REPORT_TYPES.get(report.get('reportType'), report.get('reportType'))}")
    doc.add_paragraph(f"专业：{report.get('major', '')}    电厂：{report.get('plant', '')}    年份：{report.get('year', '')}")
    for chapter in outline:
        doc.add_heading(f"{chapter['chapterNo']} {chapter['title']}", min(int(chapter.get("level", 1)), 3))
        content = contents.get(chapter["chapterId"], {})
        for paragraph in (content.get("content") or "暂无内容").splitlines():
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        for table_data in content.get("tables", []):
            if table_data.get("title"):
                doc.add_paragraph(table_data["title"])
            headers = table_data.get("headers") or []
            rows = table_data.get("rows") or []
            table = doc.add_table(rows=1, cols=max(len(headers), 1))
            table.style = "Table Grid"
            for idx, header in enumerate(headers or ["内容"]):
                table.rows[0].cells[idx].text = str(header)
            for row in rows:
                cells = table.add_row().cells
                for idx, value in enumerate(row[: len(cells)]):
                    cells[idx].text = str(value)
    doc.save(str(path))


def write_pdf(path: Path, report: dict, markdown: str) -> None:
    if canvas is None:
        raise RuntimeError("reportlab is not available")
    c = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    font_name = "Helvetica"
    # Try a Chinese-capable font on Windows.
    for candidate in [
        r"C:\\Windows\\Fonts\\simhei.ttf",
        r"C:\\Windows\\Fonts\\msyh.ttc",
        r"C:\\Windows\\Fonts\\simsun.ttc",
    ]:
        if Path(candidate).exists():
            try:
                pdfmetrics.registerFont(TTFont("ReportCN", candidate))
                font_name = "ReportCN"
                break
            except Exception:
                pass
    c.setFont(font_name, 16)
    c.drawString(50, height - 50, report.get("reportName") or "报告")
    c.setFont(font_name, 10)
    y = height - 85
    for raw_line in markdown.splitlines():
        line = raw_line.replace("#", "").strip()
        if not line:
            y -= 10
            continue
        while len(line) > 58:
            c.drawString(50, y, line[:58])
            line = line[58:]
            y -= 16
            if y < 50:
                c.showPage()
                c.setFont(font_name, 10)
                y = height - 50
        c.drawString(50, y, line)
        y -= 16
        if y < 50:
            c.showPage()
            c.setFont(font_name, 10)
            y = height - 50
    c.save()


def create_export(report_id: str, file_format: str, template_id: str | None = None) -> dict:
    if file_format not in FILE_FORMATS:
        raise ApiError(40001, "参数校验失败", data={"field": "fileFormat", "reason": "不支持的导出格式"})
    detail = report_detail(report_id)
    report = detail["report"]
    outline = detail["outline"]
    contents = {c["chapterId"]: c for c in detail["contents"]}
    export_id = store.next_id("exp")
    base_name = safe_filename(report.get("reportName") or report_id)
    file_name = f"{base_name}.{file_format}"
    rel_path = Path("exports") / f"{export_id}-{file_name}"
    abs_path = STORAGE_DIR / rel_path
    markdown = export_markdown(report, outline, contents)

    try:
        if file_format == "docx":
            write_docx(abs_path, report, outline, contents)
        elif file_format == "pdf":
            write_pdf(abs_path, report, markdown)
        elif file_format == "md":
            abs_path.write_text(markdown, encoding="utf-8")
        else:
            text = re.sub(r"^#+\\s*", "", markdown, flags=re.MULTILINE)
            abs_path.write_text(text, encoding="utf-8")
        status = "exported"
    except Exception:
        status = "exportFailed"
        traceback.print_exc()
        raise ApiError(50020, "DOCX 导出失败")

    record = {
        "exportId": export_id,
        "reportId": report_id,
        "fileName": file_name,
        "fileFormat": file_format,
        "fileSize": abs_path.stat().st_size,
        "downloadUrl": f"/api/reports/{report_id}/exports/{export_id}/download",
        "status": status,
        "createdAt": now_iso(),
        "filePath": str(rel_path).replace("\\\\", "/"),
        "templateId": template_id,
    }
    store.data["exports"][export_id] = record
    report["status"] = "exported" if status == "exported" else "exportFailed"
    report["updatedAt"] = now_iso()
    store.save()
    return public_export(record)


def public_export(export: dict) -> dict:
    return {k: v for k, v in export.items() if k != "filePath"}


def parse_multipart(body: bytes, content_type: str) -> tuple[dict, dict]:
    match = re.search(r"boundary=(.+)", content_type)
    if not match:
        return {}, {}
    boundary = match.group(1).strip().strip('"').encode("utf-8")
    fields: dict[str, str] = {}
    files: dict[str, dict] = {}
    for part in body.split(b"--" + boundary):
        part = part.strip()
        if not part or part == b"--":
            continue
        if part.endswith(b"--"):
            part = part[:-2]
        header_blob, _, content = part.partition(b"\\r\\n\\r\\n")
        if not _:
            continue
        headers = header_blob.decode("utf-8", "ignore").split("\\r\\n")
        disposition = next((h for h in headers if h.lower().startswith("content-disposition:")), "")
        name_match = re.search(r'name="([^"]+)"', disposition)
        if not name_match:
            continue
        name = name_match.group(1)
        filename_match = re.search(r'filename="([^"]*)"', disposition)
        content = content.rstrip(b"\\r\\n")
        if filename_match:
            files[name] = {"filename": filename_match.group(1), "content": content}
        else:
            fields[name] = content.decode("utf-8", "ignore")
    return fields, files


class Handler(BaseHTTPRequestHandler):
    server_version = "ReportBackend/1.0"

    def log_message(self, fmt: str, *args) -> None:
        sys.stderr.write("[%s] %s\\n" % (self.log_date_time_string(), fmt % args))

    def end_headers(self) -> None:
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Headers", "Authorization, Content-Type, Accept")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, PUT, PATCH, DELETE, OPTIONS")
        self.send_header("Access-Control-Expose-Headers", "Content-Disposition")
        super().end_headers()

    def do_OPTIONS(self) -> None:
        self.send_response(204)
        self.end_headers()

    def do_GET(self) -> None:
        self.dispatch("GET")

    def do_POST(self) -> None:
        self.dispatch("POST")

    def do_PUT(self) -> None:
        self.dispatch("PUT")

    def do_PATCH(self) -> None:
        self.dispatch("PATCH")

    def do_DELETE(self) -> None:
        self.dispatch("DELETE")

    def dispatch(self, method: str) -> None:
        try:
            parsed = urllib.parse.urlparse(self.path)
            path = parsed.path.rstrip("/") or "/"
            query = urllib.parse.parse_qs(parsed.query)
            if path == "/api/health":
                return self.send_json(success({"status": "ok", "time": now_iso()}))
            if method == "POST" and path in {"/api/reports", "/api/auth/register", "/api/auth/login", "/api/auth/logout"}:
                pass
            result = self.route(method, path, query)
            if result is not None:
                self.send_json(success(result))
        except ApiError as exc:
            self.send_json({"code": exc.code, "message": exc.message, "data": exc.data, "traceId": trace_id()}, exc.http_status)
        except Exception as exc:
            traceback.print_exc()
            self.send_json({"code": 50000, "message": "服务器内部错误", "data": {"error": str(exc)}, "traceId": trace_id()}, 500)

    def send_json(self, payload: dict, status: int = 200) -> None:
        body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def read_body(self) -> bytes:
        length = int(self.headers.get("Content-Length", "0") or "0")
        return self.rfile.read(length) if length else b""

    def read_json(self) -> dict:
        body = self.read_body()
        if not body:
            return {}
        try:
            return json.loads(body.decode("utf-8"))
        except json.JSONDecodeError:
            raise ApiError(40000, "请求参数错误")

    def current_user(self, required: bool = True, roles: set[str] | None = None) -> dict | None:
        auth = self.headers.get("Authorization", "")
        token = auth.removeprefix("Bearer ").strip() if auth.startswith("Bearer ") else ""
        session = store.data.get("sessions", {}).get(token)
        if not session:
            if required:
                raise ApiError(40100, "未登录或登录过期")
            return None
        user = store.data["users"].get(session["userId"])
        if not user or user.get("status", "enabled") != "enabled":
            raise ApiError(40100, "未登录或登录过期")
        if roles and user["role"] not in roles:
            raise ApiError(40300, "无权限")
        return user

    def require_report(self, report_id: str, user: dict) -> dict:
        report = store.data["reports"].get(report_id)
        if not report:
            raise ApiError(40400, "资源不存在")
        if not user_can_access_report(user, report):
            raise ApiError(40300, "无权限")
        return report

    def route(self, method: str, path: str, query: dict) -> dict | None:
        parts = [p for p in path.split("/") if p]
        if parts[:1] != ["api"]:
            raise ApiError(40400, "资源不存在", http_status=404)

        if parts[1:2] == ["auth"]:
            return self.route_auth(method, parts[2:])

        user = self.current_user()

        if parts[1:2] == ["reports"]:
            return self.route_reports(method, parts[2:], query, user)
        if parts[1:2] == ["templates"]:
            return self.route_templates(method, parts[2:], query, user)
        if parts[1:2] == ["materials"]:
            return self.route_materials(method, parts[2:], query, user)
        if parts[1:2] == ["admin"]:
            return self.route_admin(method, parts[2:], query, user)
        raise ApiError(40400, "资源不存在", http_status=404)

    def route_auth(self, method: str, parts: list[str]) -> dict | None:
        if method == "POST" and parts == ["register"]:
            data = self.read_json()
            username = str(data.get("username", "")).strip()
            password = str(data.get("password", "")).strip()
            if not username or not password:
                raise ApiError(40001, "参数校验失败", data={"field": "username", "reason": "用户名和密码必填"})
            if any(u["username"] == username for u in store.data["users"].values()):
                raise ApiError(40900, "状态冲突", data={"field": "username", "reason": "用户名已存在"})
            user_id = store.next_id("usr")
            user = {
                "userId": user_id,
                "username": username,
                "passwordHash": hash_password(password),
                "role": "user",
                "displayName": data.get("displayName") or username,
                "status": "enabled",
                "createdAt": now_iso(),
            }
            store.data["users"][user_id] = user
            store.save()
            return {"userId": user_id, "username": username, "role": "user"}

        if method == "POST" and parts == ["login"]:
            data = self.read_json()
            username = str(data.get("username", "")).strip()
            password = str(data.get("password", "")).strip()
            user = next((u for u in store.data["users"].values() if u["username"] == username), None)
            if not user or not verify_password(password, user.get("passwordHash", "")):
                raise ApiError(40001, "用户名或密码错误")
            if user.get("status", "enabled") != "enabled":
                raise ApiError(40300, "账号已停用")
            token = secrets.token_urlsafe(32)
            store.data["sessions"][token] = {"userId": user["userId"], "createdAt": now_iso()}
            store.save()
            return {"accessToken": token, "expiresIn": 7200, "user": public_user(user)}

        if method == "GET" and parts == ["me"]:
            return public_user(self.current_user())

        if method == "POST" and parts == ["logout"]:
            auth = self.headers.get("Authorization", "")
            token = auth.removeprefix("Bearer ").strip() if auth.startswith("Bearer ") else ""
            store.data.get("sessions", {}).pop(token, None)
            store.save()
            return None

        raise ApiError(40400, "资源不存在", http_status=404)

    def route_reports(self, method: str, parts: list[str], query: dict, user: dict) -> dict | None:
        if method == "POST" and not parts:
            data = self.read_json()
            report_type = data.get("reportType")
            if report_type not in REPORT_TYPES:
                raise ApiError(40001, "参数校验失败", data={"field": "reportType", "reason": "不支持的报告类型"})
            report_id = store.next_id("rpt")
            report_name = data.get("reportName") or f"{data.get('plant', '')}{REPORT_TYPES[report_type]}"
            report = {
                "reportId": report_id,
                "reportName": report_name,
                "reportType": report_type,
                "topic": data.get("topic") or "",
                "major": data.get("major") or "",
                "plant": data.get("plant") or "",
                "year": int(data.get("year") or datetime.now(TZ).year),
                "status": "draft",
                "createdBy": user["userId"],
                "createdAt": now_iso(),
                "generatedAt": None,
                "updatedAt": now_iso(),
            }
            store.data["reports"][report_id] = report
            store.data["outlines"][report_id] = []
            store.data["contents"][report_id] = {}
            store.save()
            return {k: report[k] for k in ("reportId", "reportName", "reportType", "status")}

        if method == "GET" and not parts:
            reports = list(store.data["reports"].values())
            if user["role"] == "user":
                reports = [r for r in reports if r.get("createdBy") == user["userId"]]
            keyword = (query.get("keyword", [""])[0] or "").strip()
            report_type = query.get("reportType", [""])[0]
            status = query.get("status", [""])[0]
            if keyword:
                reports = [r for r in reports if keyword in r.get("reportName", "") or keyword in r.get("topic", "")]
            if report_type:
                reports = [r for r in reports if r.get("reportType") == report_type]
            if status:
                reports = [r for r in reports if r.get("status") == status]
            reports.sort(key=lambda r: r.get("createdAt", ""), reverse=True)
            items = [
                {
                    "reportId": r["reportId"],
                    "reportName": r["reportName"],
                    "reportType": r["reportType"],
                    "plant": r.get("plant", ""),
                    "year": r.get("year"),
                    "status": r.get("status"),
                    "generatedAt": r.get("generatedAt"),
                    "createdAt": r.get("createdAt"),
                }
                for r in reports
            ]
            return page_items(items, int(query.get("page", [1])[0]), int(query.get("pageSize", [10])[0]))

        if not parts:
            raise ApiError(40400, "资源不存在", http_status=404)

        report_id = parts[0]
        report = self.require_report(report_id, user)

        if len(parts) == 1:
            if method == "GET":
                return report_detail(report_id)
            if method == "PUT":
                data = self.read_json()
                for field in ["reportName", "topic", "major", "plant"]:
                    if field in data:
                        report[field] = data[field]
                if "year" in data:
                    report["year"] = int(data["year"])
                report["updatedAt"] = now_iso()
                store.save()
                return report
            if method == "DELETE":
                del store.data["reports"][report_id]
                store.data.get("outlines", {}).pop(report_id, None)
                store.data.get("contents", {}).pop(report_id, None)
                for export_id, export in list(store.data.get("exports", {}).items()):
                    if export["reportId"] == report_id:
                        store.data["exports"].pop(export_id, None)
                store.save()
                return None

        if parts[1:2] == ["outline"]:
            return self.route_outline(method, report, parts[2:])
        if parts[1:2] == ["content"]:
            return self.route_content_generate(method, report)
        if parts[1:2] == ["chapters"]:
            return self.route_chapters(method, report, parts[2:])
        if parts[1:2] == ["exports"]:
            return self.route_exports(method, report, parts[2:], query)

        raise ApiError(40400, "资源不存在", http_status=404)

    def route_outline(self, method: str, report: dict, parts: list[str]) -> dict | None:
        report_id = report["reportId"]
        if method == "POST" and parts == ["generate"]:
            outline = default_outline(report)
            store.data["outlines"][report_id] = outline
            store.data["contents"].setdefault(report_id, {})
            report["status"] = "outlineGenerated"
            report["updatedAt"] = now_iso()
            store.save()
            return {"reportId": report_id, "outline": outline}
        if method == "PUT" and not parts:
            data = self.read_json()
            outline = normalize_outline(report_id, data.get("outline") or [])
            old_content = store.data.get("contents", {}).setdefault(report_id, {})
            valid_ids = {chapter["chapterId"] for chapter in outline}
            for chapter_id in list(old_content):
                if chapter_id not in valid_ids:
                    old_content.pop(chapter_id, None)
            store.data["outlines"][report_id] = outline
            report["status"] = "outlineGenerated"
            report["updatedAt"] = now_iso()
            store.save()
            return {"reportId": report_id, "outline": outline}
        raise ApiError(40400, "资源不存在", http_status=404)

    def route_content_generate(self, method: str, report: dict) -> None:
        if method != "POST":
            raise ApiError(40400, "资源不存在", http_status=404)
        data = self.read_json()
        chapter_ids = data.get("chapterIds") or []
        regenerate = bool(data.get("regenerate"))
        force_overwrite = bool(data.get("forceOverwrite"))
        outline = store.data.get("outlines", {}).get(report["reportId"], [])
        selected = [c for c in outline if not chapter_ids or c["chapterId"] in chapter_ids]
        self.send_sse_start()
        self.generate_chapters_sse(report, selected, regenerate, force_overwrite)
        return None

    def route_chapters(self, method: str, report: dict, parts: list[str]) -> dict | None:
        if len(parts) < 2:
            raise ApiError(40400, "资源不存在", http_status=404)
        chapter_id = parts[0]
        if method == "PUT" and parts[1:] == ["content"]:
            data = self.read_json()
            contents = store.data.setdefault("contents", {}).setdefault(report["reportId"], {})
            record = {
                "chapterId": chapter_id,
                "content": data.get("content") or "",
                "tables": data.get("tables") or [],
                "manualEdited": bool(data.get("manualEdited")),
                "updatedAt": now_iso(),
            }
            contents[chapter_id] = record
            for chapter in store.data.get("outlines", {}).get(report["reportId"], []):
                if chapter["chapterId"] == chapter_id:
                    chapter["status"] = "done"
            report["updatedAt"] = now_iso()
            store.save()
            return {"chapterId": chapter_id, "status": "done", "updatedAt": record["updatedAt"]}
        if method == "POST" and parts[1:] == ["regenerate"]:
            data = self.read_json()
            force_overwrite = bool(data.get("forceOverwrite"))
            outline = store.data.get("outlines", {}).get(report["reportId"], [])
            selected = [c for c in outline if c["chapterId"] == chapter_id]
            self.send_sse_start()
            self.generate_chapters_sse(report, selected, True, force_overwrite, extra_prompt=data.get("extraPrompt"))
            return None
        raise ApiError(40400, "资源不存在", http_status=404)

    def send_sse_start(self) -> None:
        self.send_response(200)
        self.send_header("Content-Type", "text/event-stream; charset=utf-8")
        self.send_header("Cache-Control", "no-cache")
        self.end_headers()

    def sse(self, event: str, data: dict) -> None:
        payload = f"event: {event}\\ndata: {json.dumps(data, ensure_ascii=False)}\\n\\n".encode("utf-8")
        self.wfile.write(payload)
        self.wfile.flush()

    def generate_chapters_sse(self, report: dict, selected: list[dict], regenerate: bool, force_overwrite: bool, extra_prompt: str | None = None) -> None:
        report_id = report["reportId"]
        contents = store.data.setdefault("contents", {}).setdefault(report_id, {})
        total = len(selected)
        completed = 0
        report["status"] = "generating"
        store.save()
        for chapter in selected:
            existing = contents.get(chapter["chapterId"])
            if existing and existing.get("manualEdited") and not force_overwrite:
                completed += 1
                self.sse("chapterDone", {"reportId": report_id, "chapterId": chapter["chapterId"], "status": "done", "skipped": True})
                continue
            chapter["status"] = "running"
            self.sse("chapterStart", {"reportId": report_id, "chapterId": chapter["chapterId"], "chapterNo": chapter["chapterNo"], "title": chapter["title"]})
            content, tables = build_chapter_content(report, chapter)
            if extra_prompt:
                content += f"\\n\\n补充要求：{extra_prompt}"
            pieces = [p for p in re.split(r"(，|。|；|\\n)", content) if p]
            buf = ""
            for piece in pieces:
                buf += piece
                self.sse("chunk", {"reportId": report_id, "chapterId": chapter["chapterId"], "contentDelta": piece})
                time.sleep(0.02)
            for table in tables:
                self.sse("table", {"reportId": report_id, "chapterId": chapter["chapterId"], "table": table})
            contents[chapter["chapterId"]] = {
                "chapterId": chapter["chapterId"],
                "content": buf,
                "tables": tables,
                "manualEdited": False,
                "updatedAt": now_iso(),
            }
            chapter["status"] = "done"
            completed += 1
            percent = int(completed * 100 / total) if total else 100
            self.sse("progress", {"reportId": report_id, "completedChapters": completed, "totalChapters": total, "percent": percent})
            self.sse("chapterDone", {"reportId": report_id, "chapterId": chapter["chapterId"], "status": "done"})
            store.save()
        report["status"] = "generated"
        report["generatedAt"] = now_iso()
        report["updatedAt"] = now_iso()
        store.save()
        self.sse("done", {"reportId": report_id, "status": "generated"})

    def route_exports(self, method: str, report: dict, parts: list[str], query: dict) -> dict | None:
        report_id = report["reportId"]
        if method == "POST" and not parts:
            data = self.read_json()
            file_format = data.get("fileFormat") or "docx"
            export = create_export(report_id, file_format, data.get("templateId"))
            return export
        if method == "GET" and not parts:
            exports = [public_export(e) for e in store.data.get("exports", {}).values() if e["reportId"] == report_id]
            file_format = query.get("fileFormat", [""])[0]
            if file_format:
                exports = [e for e in exports if e["fileFormat"] == file_format]
            exports.sort(key=lambda e: e.get("createdAt", ""), reverse=True)
            return page_items(exports, int(query.get("page", [1])[0]), int(query.get("pageSize", [10])[0]))
        if len(parts) >= 1:
            export_id = parts[0]
            export = store.data.get("exports", {}).get(export_id)
            if not export or export["reportId"] != report_id:
                raise ApiError(40400, "资源不存在")
            if method == "GET" and len(parts) == 1:
                return public_export(export)
            if method == "GET" and parts[1:] == ["download"]:
                return self.download_export(export)
        raise ApiError(40400, "资源不存在", http_status=404)

    def download_export(self, export: dict) -> None:
        path = STORAGE_DIR / export["filePath"]
        if not path.exists():
            raise ApiError(40400, "资源不存在")
        content = path.read_bytes()
        mime = mimetypes.guess_type(export["fileName"])[0] or "application/octet-stream"
        if export["fileFormat"] == "md":
            mime = "text/markdown; charset=utf-8"
        if export["fileFormat"] == "txt":
            mime = "text/plain; charset=utf-8"
        encoded_name = urllib.parse.quote(export["fileName"])
        self.send_response(200)
        self.send_header("Content-Type", mime)
        self.send_header("Content-Disposition", f"attachment; filename*=UTF-8''{encoded_name}")
        self.send_header("Content-Length", str(len(content)))
        self.end_headers()
        self.wfile.write(content)
        return None

    def route_templates(self, method: str, parts: list[str], query: dict, user: dict) -> dict | None:
        if method == "GET" and not parts:
            templates = list(store.data["templates"].values())
            report_type = query.get("reportType", [""])[0]
            if report_type:
                templates = [t for t in templates if t.get("reportType") == report_type]
            templates.sort(key=lambda t: t.get("createdAt", ""), reverse=True)
            return page_items([public_template(t) for t in templates], int(query.get("page", [1])[0]), int(query.get("pageSize", [10])[0]))
        if method == "GET" and len(parts) == 1:
            template = store.data["templates"].get(parts[0])
            if not template:
                raise ApiError(40400, "资源不存在")
            return public_template(template, detail=True)
        if method == "POST" and not parts:
            self.current_user(roles={"admin", "superAdmin"})
            fields, files = self.read_upload_or_json()
            report_type = fields.get("reportType")
            if report_type not in REPORT_TYPES:
                raise ApiError(40001, "参数校验失败", data={"field": "reportType", "reason": "不支持的报告类型"})
            template_id = store.next_id("tpl")
            file_info = files.get("file")
            file_name = file_info["filename"] if file_info else f"{template_id}.docx"
            rel = ""
            if file_info:
                rel = str(Path("uploads") / f"{template_id}-{safe_filename(file_name)}").replace("\\\\", "/")
                (STORAGE_DIR / rel).write_bytes(file_info["content"])
            template = {
                "templateId": template_id,
                "templateName": fields.get("templateName") or file_name,
                "reportType": report_type,
                "fileName": file_name,
                "filePath": rel,
                "status": "enabled",
                "structure": {"titleStyle": "Heading1", "bodyStyle": "Normal", "tableStyle": "GridTable"},
                "createdBy": user["userId"],
                "createdAt": now_iso(),
            }
            store.data["templates"][template_id] = template
            store.save()
            return public_template(template)
        if len(parts) == 1:
            self.current_user(roles={"admin", "superAdmin"})
            template = store.data["templates"].get(parts[0])
            if not template:
                raise ApiError(40400, "资源不存在")
            if method == "PUT":
                data = self.read_json()
                for field in ["templateName", "status", "structure"]:
                    if field in data:
                        template[field] = data[field]
                store.save()
                return public_template(template, detail=True)
            if method == "DELETE":
                template["status"] = "disabled"
                store.save()
                return None
        raise ApiError(40400, "资源不存在", http_status=404)

    def read_upload_or_json(self) -> tuple[dict, dict]:
        content_type = self.headers.get("Content-Type", "")
        body = self.read_body()
        if content_type.startswith("multipart/form-data"):
            return parse_multipart(body, content_type)
        if body:
            try:
                return json.loads(body.decode("utf-8")), {}
            except Exception:
                raise ApiError(40000, "请求参数错误")
        return {}, {}

    def route_materials(self, method: str, parts: list[str], query: dict, user: dict) -> dict | None:
        if method == "GET" and not parts:
            materials = list(store.data["materials"].values())
            major = query.get("major", [""])[0]
            keyword = query.get("keyword", [""])[0]
            if major:
                materials = [m for m in materials if m.get("major") == major]
            if keyword:
                materials = [m for m in materials if keyword in m.get("materialName", "")]
            materials.sort(key=lambda m: m.get("createdAt", ""), reverse=True)
            return page_items(materials, int(query.get("page", [1])[0]), int(query.get("pageSize", [10])[0]))
        if method == "POST" and not parts:
            self.current_user(roles={"admin", "superAdmin"})
            fields, files = self.read_upload_or_json()
            file_info = files.get("file")
            material_id = store.next_id("mat")
            file_name = file_info["filename"] if file_info else f"{material_id}.txt"
            rel = ""
            size = 0
            if file_info:
                rel = str(Path("uploads") / f"{material_id}-{safe_filename(file_name)}").replace("\\\\", "/")
                (STORAGE_DIR / rel).write_bytes(file_info["content"])
                size = len(file_info["content"])
            material = {
                "materialId": material_id,
                "materialName": fields.get("materialName") or file_name,
                "major": fields.get("major") or "",
                "fileName": file_name,
                "fileSize": size,
                "filePath": rel,
                "status": "enabled",
                "createdBy": user["userId"],
                "createdAt": now_iso(),
            }
            store.data["materials"][material_id] = material
            store.save()
            return material
        if len(parts) == 1:
            self.current_user(roles={"admin", "superAdmin"})
            material = store.data["materials"].get(parts[0])
            if not material:
                raise ApiError(40400, "资源不存在")
            if method == "DELETE":
                store.data["materials"].pop(parts[0], None)
                store.save()
                return None
        if len(parts) == 2 and parts[1] == "status" and method == "PATCH":
            self.current_user(roles={"admin", "superAdmin"})
            material = store.data["materials"].get(parts[0])
            if not material:
                raise ApiError(40400, "资源不存在")
            data = self.read_json()
            status = data.get("status")
            if status not in {"enabled", "disabled"}:
                raise ApiError(40001, "参数校验失败", data={"field": "status", "reason": "状态不支持"})
            material["status"] = status
            store.save()
            return material
        raise ApiError(40400, "资源不存在", http_status=404)

    def route_admin(self, method: str, parts: list[str], query: dict, user: dict) -> dict | None:
        if parts[:1] == ["model-config"]:
            self.current_user(roles={"admin", "superAdmin"})
            if method == "GET" and len(parts) == 1:
                return mask_model_config(store.data["modelConfig"])
            if method == "PUT" and len(parts) == 1:
                data = self.read_json()
                config = store.data["modelConfig"]
                for field in ["apiUrl", "modelName", "apiKey", "timeoutSeconds", "enabled"]:
                    if field in data:
                        config[field] = data[field]
                config["updatedAt"] = now_iso()
                store.save()
                return mask_model_config(config)
            if method == "POST" and parts[1:] == ["test"]:
                return {"available": True, "latencyMs": 120}
        if parts[:1] == ["users"]:
            self.current_user(roles={"superAdmin"})
            return self.route_users(method, parts[1:], query)
        raise ApiError(40400, "资源不存在", http_status=404)

    def route_users(self, method: str, parts: list[str], query: dict) -> dict | None:
        if method == "GET" and not parts:
            users = [public_user(u) for u in store.data["users"].values()]
            keyword = query.get("keyword", [""])[0]
            role = query.get("role", [""])[0]
            status = query.get("status", [""])[0]
            if keyword:
                users = [u for u in users if keyword in u["username"] or keyword in u.get("displayName", "")]
            if role:
                users = [u for u in users if u["role"] == role]
            if status:
                users = [u for u in users if u.get("status") == status]
            users.sort(key=lambda u: u.get("createdAt", ""), reverse=True)
            return page_items(users, int(query.get("page", [1])[0]), int(query.get("pageSize", [10])[0]))
        if len(parts) == 2 and method == "PATCH":
            user = store.data["users"].get(parts[0])
            if not user:
                raise ApiError(40400, "资源不存在")
            data = self.read_json()
            if parts[1] == "role":
                role = data.get("role")
                if role not in ROLES:
                    raise ApiError(40001, "参数校验失败", data={"field": "role", "reason": "角色不支持"})
                user["role"] = role
                store.save()
                return public_user(user)
            if parts[1] == "status":
                status = data.get("status")
                if status not in USER_STATUSES:
                    raise ApiError(40001, "参数校验失败", data={"field": "status", "reason": "状态不支持"})
                user["status"] = status
                store.save()
                return public_user(user)
        raise ApiError(40400, "资源不存在", http_status=404)


def public_template(template: dict, detail: bool = False) -> dict:
    data = {k: v for k, v in template.items() if k != "filePath"}
    if not detail:
        data.pop("structure", None)
    return data


def main() -> None:
    host = os.environ.get("REPORT_BACKEND_HOST", "127.0.0.1")
    port = int(os.environ.get("REPORT_BACKEND_PORT", "8000"))
    server = ThreadingHTTPServer((host, port), Handler)
    print(f"Report backend running at http://{host}:{port}")
    print("Default accounts: student/123456, admin/admin123, super/super123")
    server.serve_forever()


if __name__ == "__main__":
    main()
